from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io

COLORS = {
    "High":     RGBColor(0xC0, 0x39, 0x2B),
    "Moderate": RGBColor(0xE6, 0xA8, 0x17),
    "Low":      RGBColor(0x1E, 0x71, 0x45),
    "navy":     RGBColor(0x1B, 0x3A, 0x6B),
    "blue":     RGBColor(0x2E, 0x75, 0xB6),
}

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = COLORS["navy"]
    run.font.size = Pt(18 if level == 1 else 13)
    run.font.name = "Arial"
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    return p

def build_docx(ddr: dict, inspection_images: list = None, thermal_images: list = None) -> bytes:
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)

    # ── Cover ─────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DETAILED DIAGNOSTIC REPORT")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = COLORS["navy"]
    doc.add_paragraph()

    # ── Section 1: Property Issue Summary ─────────────────────────
    add_heading(doc, "1. Property Issue Summary")
    ps = ddr.get("property_summary", {})
    doc.add_paragraph(ps.get("overview", "Not Available"))

    if ps.get("areas"):
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = "Area"
        hdr[1].text = "Issue"
        for cell in hdr:
            set_cell_bg(cell, "1B3A6B")
            for para in cell.paragraphs:
                run = para.runs[0] if para.runs else para.add_run(cell.text)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True
        for area in ps["areas"]:
            row = table.add_row().cells
            if isinstance(area, dict):
                row[0].text = area.get("area", "")
                row[1].text = area.get("issue", "")
            else:
                row[0].text = str(area)
                row[1].text = "See area observations below"

    # ── Section 2: Area-wise Observations ─────────────────────────
    doc.add_page_break()
    add_heading(doc, "2. Area-wise Observations")

    area_image_map = {
        "hall":            [0, 1, 2, 3, 4, 5, 6],
        "bedroom":         [7, 8, 9, 10],
        "master bedroom":  [11, 12, 13, 14, 15],
        "kitchen":         [16, 17],
        "parking":         [18, 19, 20, 21],
        "common bathroom": [22, 23, 24],
    }

    thermal_image_map = {
        "hall":            [0, 1],
        "bedroom":         [2, 3],
        "master bedroom":  [4, 5],
        "kitchen":         [6, 7],
        "parking":         [8, 9],
        "common bathroom": [10, 11],
    }

    for obs in ddr.get("area_observations", []):
        area_name = obs.get("area", "")
        add_heading(doc, area_name, level=2)

        # Observation table
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        labels = ["Negative Side (Affected)", "Positive Side (Source)", "Thermal Finding"]
        keys   = ["negative_side", "positive_side", "thermal_finding"]
        for i, (label, key) in enumerate(zip(labels, keys)):
            row = table.rows[i].cells
            row[0].text = label
            set_cell_bg(row[0], "D6E4F7")
            row[0].paragraphs[0].runs[0].bold = True
            row[1].text = obs.get(key, "Not Available")
        doc.add_paragraph()

        # Inspection images for this area
        area_key = area_name.lower()
        matched_key = None
        for key in area_image_map:
            if key in area_key:
                matched_key = key
                break

        if matched_key and inspection_images:
            indices = area_image_map[matched_key]
            available = [
                inspection_images[i]
                for i in indices
                if i < len(inspection_images)
            ]
            if available:
                p = doc.add_paragraph()
                p.add_run("Site Photographs:").bold = True
                for i in range(0, min(len(available), 6), 3):
                    row_imgs = available[i:i+3]
                    img_table = doc.add_table(rows=1, cols=len(row_imgs))
                    img_table.style = 'Table Grid'
                    for col_idx, img_data in enumerate(row_imgs):
                        cell = img_table.rows[0].cells[col_idx]
                        try:
                            img_stream = io.BytesIO(img_data["bytes"])
                            paragraph = cell.paragraphs[0]
                            run = paragraph.add_run()
                            run.add_picture(img_stream, width=Inches(1.8))
                        except Exception:
                            cell.text = "Image Not Available"
                doc.add_paragraph()

        # Thermal images for this area
        thermal_matched = None
        for key in thermal_image_map:
            if key in area_key:
                thermal_matched = key
                break

        if thermal_matched and thermal_images:
            indices = thermal_image_map[thermal_matched]
            available = [
                thermal_images[i]
                for i in indices
                if i < len(thermal_images)
            ]
            if available:
                p = doc.add_paragraph()
                p.add_run("Thermal Images:").bold = True
                img_table = doc.add_table(rows=1, cols=min(len(available), 2))
                img_table.style = 'Table Grid'
                for col_idx, img_data in enumerate(available[:2]):
                    cell = img_table.rows[0].cells[col_idx]
                    try:
                        img_stream = io.BytesIO(img_data["bytes"])
                        paragraph = cell.paragraphs[0]
                        run = paragraph.add_run()
                        run.add_picture(img_stream, width=Inches(2.5))
                    except Exception:
                        cell.text = "Image Not Available"
                doc.add_paragraph()

    # ── Section 3: Root Causes ─────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "3. Probable Root Cause")
    for i, rc in enumerate(ddr.get("root_causes", []), 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {rc.get('title', '')}")
        run.bold = True
        run.font.color.rgb = COLORS["blue"]
        doc.add_paragraph(rc.get("detail", ""))
        ev = doc.add_paragraph()
        ev.add_run("Evidence: ").bold = True
        ev.add_run(rc.get("evidence", ""))
        doc.add_paragraph()

    # ── Section 4: Severity ────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "4. Severity Assessment")
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    headers = ["Area", "Severity", "Reasoning", "Recommended Action"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_bg(cell, "1B3A6B")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].runs[0].bold = True
    for sev in ddr.get("severity", []):
        row = table.add_row().cells
        row[0].text = sev.get("area", "")
        row[1].text = sev.get("level", "")
        level = sev.get("level", "")
        if level == "High":
            set_cell_bg(row[1], "C0392B")
            row[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        elif level == "Moderate":
            set_cell_bg(row[1], "E6A817")
        elif level == "Low":
            set_cell_bg(row[1], "1E7145")
            row[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        row[2].text = sev.get("reasoning", "")
        row[3].text = sev.get("action", "")

    # ── Section 5: Recommended Actions ────────────────────────────
    doc.add_page_break()
    add_heading(doc, "5. Recommended Actions")
    actions = ddr.get("recommended_actions", {})
    for phase, label in [
        ("immediate",  "Phase 1 — Immediate (Within 7 Days)"),
        ("short_term", "Phase 2 — Short-Term (Within 30 Days)"),
        ("long_term",  "Phase 3 — Long-Term (After Drying)")
    ]:
        add_heading(doc, label, level=2)
        for item in actions.get(phase, []):
            doc.add_paragraph(item, style='List Bullet')

    # ── Section 6: Additional Notes ───────────────────────────────
    doc.add_page_break()
    add_heading(doc, "6. Additional Notes")
    for note in ddr.get("additional_notes", []):
        doc.add_paragraph(note, style='List Bullet')

    # ── Section 7: Missing or Unclear Information ─────────────────
    doc.add_page_break()
    add_heading(doc, "7. Missing or Unclear Information")
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    for i, h in enumerate(["Information Item", "Status"]):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_bg(cell, "1B3A6B")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].runs[0].bold = True
    for item in ddr.get("missing_info", []):
        row = table.add_row().cells
        row[0].text = item.get("item", "")
        row[1].text = item.get("status", "Not Available")

    # Save to bytes buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()